"""
报告生成引擎
- 占位符解析与填充
- CSV 模板生成与解析
- Word 文档导出
- 报告对比
"""
import csv
import re
from datetime import datetime
from io import BytesIO, StringIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ReportEngine:
    """报告生成引擎"""

    # ── 占位符处理 ──

    @staticmethod
    def extract_placeholders(template: dict) -> list[str]:
        """从模板中提取所有占位符"""
        placeholders = set()
        for section in template.get("sections", []):
            content = section.get("content", "")
            found = re.findall(r"\{\{(.+?)\}\}", content)
            placeholders.update(found)
        return sorted(placeholders)

    @staticmethod
    def fill_placeholders(template: dict, data: dict) -> dict:
        """用数据填充模板中的占位符，返回填充后的报告"""
        sections = []
        for sec in template.get("sections", []):
            content = sec.get("content", "")
            for key, val in data.items():
                val_str = str(val) if val is not None else f"[未填写: {key}]"
                content = content.replace("{{" + key + "}}", val_str)
            sections.append({"title": sec["title"], "content": content})
        return {"sections": sections}

    # ── CSV 模板 ──

    @staticmethod
    def generate_csv_template(placeholders: list[str]) -> BytesIO:
        """生成 CSV 模板文件供用户下载填写"""
        output = StringIO()
        writer = csv.writer(output)
        writer.writerow(placeholders)  # 表头 = 占位符名
        writer.writerow([""] * len(placeholders))  # 示例空行
        buf = BytesIO(output.getvalue().encode("utf-8-sig"))
        buf.seek(0)
        return buf

    @staticmethod
    def parse_csv(csv_bytes: bytes) -> dict:
        """解析用户上传的 CSV 文件，返回 {占位符: 值} 字典"""
        content = csv_bytes.decode("utf-8-sig")
        reader = csv.reader(StringIO(content))
        rows = list(reader)
        if not rows:
            raise ValueError("CSV 文件为空")

        headers = [h.strip() for h in rows[0]]
        values = rows[1] if len(rows) > 1 else [""] * len(headers)

        result = {}
        for i, header in enumerate(headers):
            val = values[i].strip() if i < len(values) else ""
            result[header] = val

        return result

    # ── Word 导出 ──

    @staticmethod
    def export_word(report: dict) -> BytesIO:
        """将报告导出为 Word 文档"""
        doc = Document()

        # 标题
        title = doc.add_heading(report.get("title", "报告"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        info_parts = []
        if report.get("report_no"):
            info_parts.append(f"报告编号：{report['report_no']}")
        if report.get("author"):
            info_parts.append(f"编制人：{report['author']}")
        if report.get("date"):
            info_parts.append(f"日期：{report['date']}")
        if info_parts:
            info_text = "　　".join(info_parts)
            info_p = doc.add_paragraph(info_text)
            info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in info_p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        # 章节内容
        for sec in report.get("sections", []):
            doc.add_heading(sec["title"], level=1)
            content = sec.get("content", "")
            if content:
                for line in content.split("\n"):
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.size = Pt(11)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ── 报告对比 ──

    @staticmethod
    def compare_reports(report_a: dict, report_b: dict) -> list[dict]:
        """
        对比两份报告的差异，返回逐章节对比结果
        每个 section 返回 type: 'same' | 'modified' | 'added' | 'removed'
        """
        sections_a = {s["title"]: s["content"] for s in report_a.get("sections", [])}
        sections_b = {s["title"]: s["content"] for s in report_b.get("sections", [])}

        all_titles = list(dict.fromkeys(
            list(sections_a.keys()) + list(sections_b.keys())
        ))

        results = []
        for title in all_titles:
            content_a = sections_a.get(title)
            content_b = sections_b.get(title)

            if content_a is not None and content_b is not None:
                if content_a == content_b:
                    results.append({"title": title, "type": "same", "content": content_a})
                else:
                    results.append({
                        "title": title,
                        "type": "modified",
                        "content_a": content_a,
                        "content_b": content_b,
                        "diff": ReportEngine._line_diff(content_a, content_b),
                    })
            elif content_a is not None:
                results.append({"title": title, "type": "removed", "content_a": content_a})
            else:
                results.append({"title": title, "type": "added", "content_b": content_b})

        return results

    @staticmethod
    def _line_diff(text_a: str, text_b: str) -> list[dict]:
        """逐行对比两个文本，标记增/删/相同"""
        lines_a = text_a.split("\n")
        lines_b = text_b.split("\n")

        # 简单 LCS diff
        result = []
        i, j = 0, 0
        while i < len(lines_a) or j < len(lines_b):
            if i < len(lines_a) and j < len(lines_b) and lines_a[i] == lines_b[j]:
                result.append({"type": "same", "text": lines_a[i]})
                i += 1
                j += 1
            elif i < len(lines_a) and (j >= len(lines_b) or lines_a[i] not in lines_b[j:]):
                result.append({"type": "removed", "text": lines_a[i]})
                i += 1
            elif j < len(lines_b):
                result.append({"type": "added", "text": lines_b[j]})
                j += 1
            else:
                break

        return result
